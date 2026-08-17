#!/usr/bin/env python3
"""Fact Lens 공식 양식 결과물의 구조와 문구를 회귀 검사합니다."""
from hashlib import sha256
import json
from pathlib import Path

from docx import Document

PROJECT = Path(__file__).resolve().parents[2]
KCF_ROOT = PROJECT.parent
TEMPLATE = KCF_ROOT / "공식대회_원본" / "양식"
OUTPUT = KCF_ROOT / "Fact_Lens_제출용_최종" / "필수"


def all_text(doc) -> str:
    items = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                items.append(cell.text)
    return "\n".join(items)


def margins(doc) -> list[int | None]:
    section = doc.sections[0]
    return [
        section.page_width,
        section.page_height,
        section.top_margin,
        section.bottom_margin,
        section.left_margin,
        section.right_margin,
    ]


def main() -> None:
    form1 = Document(str(OUTPUT / "서식1_Fact_Lens_작품요약서.docx"))
    form2 = Document(str(OUTPUT / "서식2_Fact_Lens_아이디어기획서.docx"))
    template2 = Document(str(TEMPLATE / "서식 2 해커톤 작품설명서(아이디어 기획서).docx"))
    form1_table = form1.tables[0]
    form2_text = all_text(form2)

    checks = {
        "form1_table_shape": [len(form1.tables), len(form1_table.rows), len(form1_table.columns)] == [1, 8, 3],
        "form1_project_name": "Fact Lens" in form1_table.rows[1].cells[1].text,
        "form1_team_blank": not form1_table.rows[0].cells[1].text.strip(),
        "form1_people_blank": all(
            not form1_table.rows[row].cells[column].text.strip()
            for row in (3, 4, 5)
            for column in (1, 2)
        ),
        "form2_official_intro_preserved": (
            [paragraph.text for paragraph in form2.paragraphs[:7]]
            == [paragraph.text for paragraph in template2.paragraphs[:7]]
        ),
        "form2_page_and_margins_preserved": margins(form2) == margins(template2),
        "form2_chosen_branch": "1. (계획 수립 단계인 경우) 순수 SW 작품 설명서" in form2_text,
        "form2_other_branches_removed": all(
            text not in form2_text
            for text in (
                "2. (계획 수립 단계인 경우) 순수 연구",
                "3. (시제품 완성 시",
                "4. (연구 완료 시)",
            )
        ),
        "form2_four_sections": all(
            text in form2_text
            for text in (
                "Ⅰ",
                "Ⅱ",
                "Ⅲ",
                "Ⅳ",
                "주제",
                "작품 설계 및 제작",
                "작품 테스트 계획",
                "제작 현황 및 향후 발전 방안",
            )
        ),
        "current_model": "Gemini 3.5 Flash-Lite" in form2_text,
        "search_grounding": "Google Search" in form2_text,
        "honest_test_state": "5개 파일, 6개 테스트 통과" in form2_text and "HTTP 429" in form2_text,
        "sources_present": "붙임. 출처" in form2_text and "ai.google.dev" in form2_text,
        "old_terms_absent": all(
            text not in form2_text
            for text in (
                "OpenAI",
                "KoELECTRA",
                "Gemma 4 31B",
                "gemma-4-31b-it",
                "KOSIS",
                "Google Search API Key",
                "Search Engine ID",
            )
        ),
        "placeholder_guidance_removed": "작품를 추진하게 된 배경 혹은 동기" not in form2_text,
    }

    application_source = (TEMPLATE / "해커톤 참가신청서.docx").read_bytes()
    application_copy = (OUTPUT / "참가신청서_Fact_Lens_작성필요.docx").read_bytes()
    checks["application_exact_copy"] = sha256(application_source).digest() == sha256(application_copy).digest()

    print(json.dumps(checks, ensure_ascii=False, indent=2))
    print(f"all_checks={all(checks.values())}")
    print(f"form2_table_count={len(form2.tables)} paragraph_count={len(form2.paragraphs)}")
    if not all(checks.values()):
        raise SystemExit(1)


if __name__ == "__main__":
    main()
