#!/usr/bin/env python3
"""Fact Lens 한국코드페어 공식 양식 제출 문서를 생성합니다.

공식 DOCX 원본은 읽기만 하고, 서식 1의 기존 표와 서식 2의 안내문·여백·
첫 번째 계획 수립 단계 순수 SW 갈래·로마숫자 대단원 표를 유지합니다.
개인정보와 팀 정보는 임의로 만들지 않고 빈칸으로 둡니다.
"""
from __future__ import annotations

from pathlib import Path
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

PROJECT = Path(__file__).resolve().parents[2]
KCF_ROOT = PROJECT.parent
OFFICIAL = KCF_ROOT / "공식대회_원본" / "양식"
OUTPUT_ROOT = KCF_ROOT / "Fact_Lens_제출용_최종"
REQUIRED = OUTPUT_ROOT / "필수"
WORK = OUTPUT_ROOT / "작업기록"

FORM1_TEMPLATE = OFFICIAL / "서식 1 해커톤 작품 요약서.docx"
FORM2_TEMPLATE = OFFICIAL / "서식 2 해커톤 작품설명서(아이디어 기획서).docx"
APPLICATION_TEMPLATE = OFFICIAL / "해커톤 참가신청서.docx"
FORM1_OUT = REQUIRED / "서식1_Fact_Lens_작품요약서.docx"
FORM2_OUT = REQUIRED / "서식2_Fact_Lens_아이디어기획서.docx"
APPLICATION_OUT = REQUIRED / "참가신청서_Fact_Lens_작성필요.docx"

FONT = "맑은 고딕"
BLACK = RGBColor(0, 0, 0)
GRAY_FILL = "E7E6E6"
WHITE_FILL = "FFFFFF"


def set_run(run, *, size: float = 10.5, bold: bool | None = None) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    node = OxmlElement("w:shd")
    node.set(qn("w:val"), "clear")
    node.set(qn("w:fill"), fill)
    tc_pr.append(node)


def set_cell_margins(cell, top: int = 80, start: int = 90, bottom: int = 80, end: int = 90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = "7F7F7F", size: str = "4") -> None:
    tbl_pr = table._tbl.tblPr
    old = tbl_pr.find(qn("w:tblBorders"))
    if old is not None:
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)
        borders.append(node)
    tbl_pr.append(borders)


def no_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 10,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(text)
    set_run(run, size=size, bold=bold)


def add_paragraph_before(doc, anchor, text: str = "", *, kind: str = "body",
                         keep_with_next: bool = False):
    paragraph = doc.add_paragraph()
    anchor.addprevious(paragraph._p)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = keep_with_next

    if kind == "heading":
        size, bold, before, after = 12, True, 8, 4
    elif kind == "subheading":
        size, bold, before, after = 10.8, True, 5, 2
    elif kind == "caption":
        size, bold, before, after = 10, False, 1, 4
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif kind == "note":
        size, bold, before, after = 10, False, 1, 3
    else:
        size, bold, before, after = 10.5, False, 0, 3

    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    run = paragraph.add_run(text)
    set_run(run, size=size, bold=bold)
    return paragraph


def add_bullets_before(doc, anchor, items: list[str]) -> None:
    for item in items:
        paragraph = add_paragraph_before(doc, anchor, f"• {item}", kind="body")
        paragraph.paragraph_format.left_indent = Cm(0.45)
        paragraph.paragraph_format.first_line_indent = Cm(-0.25)


def add_table_before(doc, anchor, headers: list[str], rows: list[list[str]],
                     widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    anchor.addprevious(table._tbl)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False if widths else True
    set_table_borders(table)

    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, size=10,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[index], GRAY_FILL)
        if widths:
            table.rows[0].cells[index].width = Cm(widths[index])
    no_split(table.rows[0])
    repeat_header(table.rows[0])

    for row_values in rows:
        row = table.add_row()
        no_split(row)
        for index, value in enumerate(row_values):
            set_cell_text(row.cells[index], str(value), size=10)
            set_cell_shading(row.cells[index], WHITE_FILL)
            if widths:
                row.cells[index].width = Cm(widths[index])
    return table


def element_text(element) -> str:
    return "".join(element.xpath(".//w:t/text()"))


def find_first_four_section_tables(doc) -> list:
    found = []
    for element in doc._element.body:
        if element.tag != qn("w:tbl"):
            continue
        text = element_text(element)
        if any(text.startswith(roman) for roman in ("Ⅰ", "Ⅱ", "Ⅲ", "Ⅳ")):
            found.append(element)
            if len(found) == 4:
                return found
    raise RuntimeError("공식 서식 2의 첫 번째 순수 SW 대단원 표 4개를 찾지 못했습니다.")


def clear_between(start, end) -> None:
    current = start.getnext()
    while current is not None and current is not end:
        following = current.getnext()
        current.getparent().remove(current)
        current = following


def format_official_intro(doc) -> None:
    for paragraph in doc.paragraphs[:10]:
        for run in paragraph.runs:
            run.font.color.rgb = BLACK
    choice = next((p for p in doc.paragraphs if p.text.startswith("1. (계획 수립 단계")), None)
    if choice is not None:
        clear_paragraph(choice)
        run = choice.add_run("1. (계획 수립 단계인 경우) 순수 SW 작품 설명서")
        set_run(run, size=11, bold=True)
        choice.paragraph_format.space_before = Pt(6)
        choice.paragraph_format.space_after = Pt(7)
        choice.paragraph_format.keep_with_next = True


def fill_form1() -> None:
    doc = Document(str(FORM1_TEMPLATE))
    table = doc.tables[0]

    # 팀명과 개인정보는 사용자가 직접 입력하도록 빈칸을 유지합니다.
    set_cell_text(table.rows[0].cells[1], "", size=11)
    set_cell_text(table.rows[1].cells[1], "Fact Lens(팩트 렌즈)", bold=True, size=11)
    for row_index in (3, 4, 5):
        set_cell_text(table.rows[row_index].cells[1], "", size=10.5)
        set_cell_text(table.rows[row_index].cells[2], "", size=10.5)

    summary_cell = table.rows[7].cells[0]
    summary_cell.text = ""
    set_cell_margins(summary_cell, top=110, start=130, bottom=100, end=130)
    sections = [
        (
            "1. 작품 주제",
            "뉴스 기사 속 주장과 표현을 기사 화면에서 바로 분석하고, 검색 근거와 함께 확인하도록 돕는 Chrome 확장 프로그램입니다. 대회 주제인 ‘신뢰할 수 있는 AI와 인간이 함께 만들어가는 안전하고 포용적인 미래 사회’에 맞춰 AI가 결론을 대신 내리지 않고 사용자가 출처를 보고 판단하도록 설계했습니다.",
        ),
        (
            "2. 작품의 설명",
            "사용자가 기사 상단의 ‘Fact Lens로 팩트체크’ 버튼을 누르면 본문을 추출합니다. Gemini Interactions API로 감정, 검증 가능한 주장, 사실·의견 비율, 누락 맥락과 프레임을 분석한 뒤, built-in Google Search로 각 주장을 검증합니다. 검색 실행 단계와 완료 상태를 확인하고 인용 URL을 저장합니다. 검증됨·미확인·거짓 상태는 기사 문장에 초록·주황·빨강 형광펜으로 표시하며, 클릭하면 설명을 보여 줍니다. 상세 결과는 확장 프로그램 popup에서 확인합니다.",
        ),
        (
            "3. 기존 해법과의 차이점",
            "기사 문장을 다른 서비스에 복사하지 않고 읽던 화면에서 분석을 시작합니다. 요약 답변만 보여 주는 대신 주장과 원문 위치를 연결하고, 검색이 실제로 실행되지 않았거나 응답이 완료되지 않으면 성공으로 처리하지 않습니다. 감정 표현과 기사 프레임도 함께 보여 주어 사실 확인과 읽기 관점을 한 흐름에서 살펴볼 수 있습니다.",
        ),
        (
            "4. 결과",
            "Vite·React·TypeScript·Manifest V3 기반 시제품을 제작했습니다. 현재 단위·통합 테스트 5개 파일의 6개 테스트와 배포 빌드가 통과합니다. API 할당량 초과나 분석 실패 시 로딩 효과를 제거하고 버튼을 다시 활성화하며 구체적인 오류를 표시합니다. 다음 단계에서는 다양한 언론사 기사, 출처 품질, 색각 접근성과 할당량 복구를 추가 검증합니다.",
        ),
    ]
    for index, (heading, body) in enumerate(sections):
        paragraph = summary_cell.paragraphs[0] if index == 0 else summary_cell.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0 if index == 0 else 3)
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.05
        h = paragraph.add_run(heading + "\n")
        set_run(h, size=10.5, bold=True)
        r = paragraph.add_run(body)
        set_run(r, size=10.3, bold=False)

    # 빈 양식의 큰 최소 높이가 내용과 결합해 2쪽으로 넘어가는 것을 막습니다.
    tr_pr = table.rows[7]._tr.get_or_add_trPr()
    for height in tr_pr.findall(qn("w:trHeight")):
        tr_pr.remove(height)

    # 표 뒤의 원본 빈 문단이 채운 표 높이와 결합하면 LibreOffice에서 빈 2쪽을 만듭니다.
    # 화면에 보이는 공식 요소가 아니므로 마지막 빈 문단만 제거합니다.
    trailing = doc.paragraphs[-1]
    if not trailing.text.strip():
        trailing._element.getparent().remove(trailing._element)

    doc.save(str(FORM1_OUT))


def fill_form2() -> None:
    doc = Document(str(FORM2_TEMPLATE))
    format_official_intro(doc)
    section_1, section_2, section_3, section_4 = find_first_four_section_tables(doc)
    sect_pr = doc._element.body.find(qn("w:sectPr"))
    if sect_pr is None:
        raise RuntimeError("서식 2의 섹션 설정을 찾지 못했습니다.")

    clear_between(section_1, section_2)
    clear_between(section_2, section_3)
    clear_between(section_3, section_4)
    clear_between(section_4, sect_pr)

    # Ⅰ. 주제
    add_paragraph_before(doc, section_2, "1. 개발 배경 및 필요성", kind="heading", keep_with_next=True)
    add_paragraph_before(
        doc, section_2,
        "뉴스 기사를 읽다가 사실 여부가 궁금한 문장을 만나면 검색창, 팩트체크 사이트, 생성형 AI 화면을 오가게 됩니다. 이 과정에서 원문 위치와 검색 결과의 관계가 끊기고, 어떤 문장이 검증 대상이었는지 다시 찾아야 합니다. 일반 생성형 AI의 짧은 답변만 읽으면 실제 검색 여부와 인용 근거를 놓치기도 합니다.",
    )
    add_paragraph_before(
        doc, section_2,
        "Fact Lens는 이 문제를 ‘기사 맥락을 유지한 채 주장 단위로 근거를 확인하기 어렵다’로 정의합니다. 사용자가 기사 페이지에서 직접 분석을 시작하고, AI는 감정·주장·편향 분석과 검색 결과 정리를 맡습니다. 최종 판단은 사용자가 인용 출처를 열어 보고 내립니다.",
    )
    add_table_before(doc, section_2,
        ["문제", "기존 흐름", "Fact Lens의 대응"],
        [
            ["맥락 단절", "문장을 복사해 다른 탭에서 검색", "기사 문장 위에 검증 상태를 직접 표시"],
            ["검색 여부 불명확", "답변이 검색 결과인지 사전지식인지 구분하기 어려움", "google_search_call 단계가 없으면 팩트체크 실패 처리"],
            ["근거 접근성", "출처를 다시 찾아야 함", "설명과 인용 링크를 tooltip·popup에 함께 제공"],
            ["표현과 프레임", "사실 확인과 별도 분석", "감정, 사실·의견 비율, 누락 맥락을 같은 보고서에 표시"],
        ], widths=[3.0, 6.0, 7.0])

    add_paragraph_before(doc, section_2, "2. 주제(목적)", kind="heading", keep_with_next=True)
    add_paragraph_before(
        doc, section_2,
        "작품의 목적은 뉴스 독자가 읽던 화면을 떠나지 않고 주장과 근거의 연결을 확인하도록 돕는 것입니다. ‘신뢰할 수 있는 AI’는 점수 하나로 단정하지 않고, 검색 실행 확인, 인용 출처, 미확인 상태, 실패 이유를 공개하는 방식으로 다룹니다.",
    )
    add_bullets_before(doc, section_2, [
        "기사 본문을 사이트별 선택자, JSON-LD, 범용 article 요소 순서로 추출합니다.",
        "한 번의 분석 호출로 감정, 최대 5개 주장, 사실·의견 비율, 누락 맥락과 기사 프레임을 구조화합니다.",
        "두 번째 호출에서 built-in Google Search를 사용하고, 검색 단계와 인용 URL이 있는지 확인합니다.",
        "검증됨·미확인·거짓 상태를 기사 원문에 표시하고 상세 설명과 출처를 제공합니다.",
        "API 오류, 미완료 응답, 할당량 초과를 숨기지 않고 사용자에게 구체적으로 알립니다.",
    ])
    add_table_before(doc, section_2,
        ["대회 주제 키워드", "작품에서의 구현"],
        [
            ["신뢰할 수 있는 AI", "검색 단계·완료 상태·인용 링크·미확인 상태를 공개"],
            ["인간과 함께", "분석 시작과 출처 판단을 사용자가 담당하고 AI는 보조"],
            ["안전한 미래", "실패를 성공으로 위장하지 않고 자동 차단·검열 도구로 사용하지 않음"],
            ["포용적인 미래", "기사 화면 안에서 색상·문장·설명으로 결과를 단계적으로 제시하고 접근성 개선을 계획"],
        ], widths=[4.2, 11.8])

    add_paragraph_before(doc, section_2, "3. 유사 제품(연구) 및 차별점", kind="heading", keep_with_next=True)
    add_paragraph_before(
        doc, section_2,
        "유사 해법은 독립형 팩트체크 서비스, 일반 생성형 AI 채팅, 기사 요약·편향 분석 도구로 나눌 수 있습니다. Fact Lens는 각각의 기능을 모두 대체한다고 주장하지 않습니다. 기사 원문과 주장, 검색 근거를 한 화면에서 연결하는 읽기 흐름에 집중합니다.",
    )
    add_table_before(doc, section_2,
        ["유사 해법", "장점", "Fact Lens와의 차이"],
        [
            ["독립형 팩트체크 서비스", "전문 검증 결과를 제공", "검증된 주제가 제한될 수 있어 기사 안의 여러 문장을 즉시 연결하기 어려움"],
            ["일반 생성형 AI 채팅", "자유롭게 질문 가능", "복사·붙여넣기가 필요하고 원문 위치와 답변 근거가 분리됨"],
            ["기사 요약·편향 분석 도구", "긴 내용을 빠르게 정리", "주장별 검색 실행과 인용을 원문 형광펜에 연결하는 흐름이 핵심은 아님"],
            ["Fact Lens", "기사 안에서 주장·출처·표현 분석을 연결", "전문기관의 최종 판정이 아닌 읽기 보조 도구이며 사용자가 출처를 확인해야 함"],
        ], widths=[3.5, 5.0, 7.5])

    # Ⅱ. 작품 설계 및 제작
    add_paragraph_before(doc, section_3, "1. 작품 설계", kind="heading", keep_with_next=True)
    add_paragraph_before(doc, section_3, "■ 설계 개념", kind="subheading", keep_with_next=True)
    add_bullets_before(doc, section_3, [
        "맥락 유지: 기사 페이지 안에서 분석을 시작하고 결과를 원문 문장에 연결합니다.",
        "근거 우선: verified·false 판정은 Google Search 실행이 확인된 경우에만 허용합니다.",
        "불확실성 공개: 근거 부족이나 결과 누락은 unverified 또는 오류로 표시합니다.",
        "최소 권한: Manifest 권한은 activeTab, storage와 Gemini API 호스트로 제한합니다.",
        "인간 판단: 자동 차단·검열을 하지 않고 사용자가 출처를 열어 판단하도록 합니다.",
    ])

    add_paragraph_before(doc, section_3, "■ 작동 원리", kind="subheading", keep_with_next=True)
    add_table_before(doc, section_3,
        ["단계", "처리 내용", "사용자에게 보이는 결과"],
        [
            ["1. 시작", "기사 상단 버튼 클릭", "본문 전체 Shimmer와 진행 상태"],
            ["2. 추출", "JSON-LD → 사이트 선택자 → 범용 선택자 → 본문 fallback", "추출 실패 시 즉시 오류"],
            ["3. 기사 분석", "Gemini 3.5 Flash-Lite가 감정·주장·편향을 JSON으로 반환", "분석 중 표시"],
            ["4. 팩트체크", "built-in Google Search로 최대 5개 주장 일괄 검증", "검색 미실행 시 실패"],
            ["5. 검증", "HTTP 상태, completed 상태, model_output, citation 확인", "불완전 응답을 성공으로 처리하지 않음"],
            ["6. 원문 표시", "정규화·슬라이딩 윈도우·문장 유사도로 주장 위치 탐색", "초록·주황·빨강 형광펜과 tooltip"],
            ["7. 상세 보고서", "결과를 chrome.storage.local에 저장", "popup에서 신뢰도·감정·편향·출처 확인"],
        ], widths=[2.0, 8.3, 5.7])

    add_paragraph_before(doc, section_3, "■ 소프트웨어 구조", kind="subheading", keep_with_next=True)
    add_table_before(doc, section_3,
        ["구성 요소", "역할", "구현 근거"],
        [
            ["Content Script", "기사 추출, 버튼·Shimmer·형광펜·tooltip", "src/content/index.ts"],
            ["Service Worker", "설정 확인, 분석 실행, 상태·결과 저장, 오류 전달", "src/background/index.ts"],
            ["Analysis Pipeline", "기사 분석 → 검색 팩트체크 → 요약 점수 계산", "src/utils/analysisPipeline.ts"],
            ["Interactions Client", "API header, completed 상태, model_output·JSON 파싱", "src/utils/geminiInteractions.ts"],
            ["Popup", "감정·팩트체크·편향 탭과 안전한 출처 링크", "src/popup/App.tsx"],
            ["Options", "Gemini API 키 입력·저장", "public/options.html·options.js"],
        ], widths=[3.3, 7.6, 5.1])

    add_paragraph_before(doc, section_3, "■ 핵심 기술 선택", kind="subheading", keep_with_next=True)
    add_table_before(doc, section_3,
        ["항목", "선택", "이유"],
        [
            ["확장 구조", "Chrome Manifest V3", "기사 DOM과 background API 호출을 분리"],
            ["화면", "React 18, TypeScript, Tailwind CSS, ShadCN UI", "popup 상태와 결과 탭을 컴포넌트로 관리"],
            ["빌드", "Vite 6", "popup·service worker·content script를 다중 entry로 묶음"],
            ["AI", "Gemini 3.5 Flash-Lite, Interactions API", "구조화 출력과 typed step 기반 검증"],
            ["검색", "built-in Google Search", "별도 검색 키 없이 검색 호출과 URL citation을 한 응답에서 처리"],
            ["개인정보", "store:false, 브라우저 저장소", "상호작용 서버 저장을 끄고 결과는 확장 로컬 저장소에 보관"],
        ], widths=[3.0, 5.3, 7.7])
    add_paragraph_before(
        doc, section_3,
        "Interactions API는 Beta이므로 schema 변경 가능성을 위험 요소로 기록합니다. API 키는 비밀번호 입력 필드와 Chrome 동기화 저장소로 관리하며 문서·코드·로그에 값 자체를 기록하지 않습니다. 기사 본문은 분석을 위해 Gemini API로 전송되므로, 개인정보가 포함된 문서에는 사용하지 않도록 안내합니다.",
        kind="note",
    )

    add_paragraph_before(doc, section_3, "2. 제작 일정", kind="heading", keep_with_next=True)
    add_table_before(doc, section_3,
        ["단계", "주요 작업", "상태"],
        [
            ["1단계", "Manifest V3 골격, 기사 추출, 분석 버튼, 기본 popup", "완료"],
            ["2단계", "Interactions API 통합, 감정·주장·편향 분석, Google Search 팩트체크", "완료"],
            ["3단계", "형광펜 매칭, tooltip, 인용 링크, 실패 상태·storage 관리", "완료"],
            ["4단계", "단위·통합 테스트, production build, 실제 브라우저 재현", "진행"],
            ["5단계", "다양한 언론사 기사·출처 품질·색각 접근성·쿼터 복구 검증", "예정"],
        ], widths=[2.2, 10.8, 3.0])

    # Ⅲ. 작품 테스트 계획
    add_paragraph_before(doc, section_4, "1. 테스트 목적", kind="heading", keep_with_next=True)
    add_paragraph_before(
        doc, section_4,
        "빌드 성공만 확인하지 않고, 기사 본문이 추출된 뒤 실제 검색 단계와 인용이 저장되고 UI에 나타나는지 검증합니다. 실패 응답을 빈 결과나 성공 화면으로 숨기지 않는 것도 같은 비중으로 확인합니다.",
    )

    add_paragraph_before(doc, section_4, "2. 테스트 방법", kind="heading", keep_with_next=True)
    add_table_before(doc, section_4,
        ["구분", "방법", "통과 조건"],
        [
            ["Interactions client", "Vitest로 HTTP 200이지만 incomplete인 응답을 주입", "completed가 아니면 명시적 오류"],
            ["기사 분석", "구조화된 감정·주장·편향 mock 응답", "필수 필드 정규화, 빈 감정 거부"],
            ["팩트체크", "google_search_call·result·URL citation mock 응답", "검색 단계가 없으면 실패, URL 저장"],
            ["전체 pipeline", "두 Interaction을 순서대로 mock", "감정·주장·팩트체크·편향·요약이 비어 있지 않음"],
            ["출처 링크", "React 정적 렌더 테스트", "http·https만 새 탭 링크, noopener·noreferrer 적용"],
            ["브라우저 E2E", "로컬 뉴스 fixture와 실제 MV3 service worker를 CDP로 실행", "버튼→API→storage→형광펜→popup 흐름 및 오류 복구"],
            ["정적 검증", "TypeScript 검사와 Vite production build", "컴파일 오류 0, dist 생성"],
        ], widths=[3.0, 7.8, 5.2])

    add_paragraph_before(doc, section_4, "3. 예상 결과", kind="heading", keep_with_next=True)
    add_bullets_before(doc, section_4, [
        "지원되는 기사에서는 본문을 추출하고 최대 5개 검증 가능한 주장을 반환합니다.",
        "검색이 실행된 경우 각 주장은 검증됨·미확인·거짓 중 하나와 설명·출처 URL을 가집니다.",
        "원문과 일치하는 주장은 상태별 형광펜으로 표시되고 popup에 동일한 결과가 나타납니다.",
        "할당량 초과, 네트워크 오류, 미완료·잘린 JSON 응답은 성공으로 저장되지 않습니다.",
        "AI 판단은 틀릴 수 있으므로 사용자는 인용 링크와 원문 맥락을 함께 확인합니다.",
    ])

    add_paragraph_before(doc, section_4, "4. 기대효과", kind="heading", keep_with_next=True)
    add_table_before(doc, section_4,
        ["대상", "기대효과"],
        [
            ["뉴스 독자", "검색 탭을 오가는 시간을 줄이고 주장과 근거의 연결을 기사 안에서 확인"],
            ["청소년·교사", "같은 기사에서 감정 표현, 사실·의견 구분, 출처 읽기를 함께 연습하는 미디어 리터러시 자료"],
            ["사회", "AI 답변을 정답으로 소비하기보다 검색 여부와 출처를 확인하는 사용 습관을 확산"],
        ], widths=[3.5, 12.5])
    add_paragraph_before(
        doc, section_4,
        "기대효과는 현재 시제품에서 관찰한 사용자 성과가 아니라 작품이 완성되고 사용자 검증을 거쳤을 때 기대하는 변화입니다. 정확도와 읽기 시간 개선은 별도 사용자 평가로 측정해야 합니다.",
        kind="note",
    )
    section_break = add_paragraph_before(doc, section_4, "")
    section_break.add_run().add_break(WD_BREAK.PAGE)

    # Ⅳ. 제작 현황 및 향후 발전 방안
    add_paragraph_before(doc, sect_pr, "1. 제작 현황", kind="heading", keep_with_next=True)
    add_paragraph_before(
        doc, sect_pr,
        "2026년 8월 16일 기준으로 Vite·React·TypeScript·Manifest V3 시제품을 구현했습니다. 기사 분석, 주장 추출, 편향 분석, Google Search 팩트체크, 결과 저장, 형광펜·tooltip, popup 보고서와 설정 화면이 코드에 연결되어 있습니다.",
    )
    add_table_before(doc, sect_pr,
        ["검증 항목", "확인 결과", "해석"],
        [
            ["단위·통합 테스트", "5개 파일, 6개 테스트 통과", "Interaction 상태·JSON·검색 step·인용·pipeline·링크 검증"],
            ["Production build", "TypeScript와 Vite build 통과", "dist 확장 프로그램 생성"],
            ["실제 기사 분석 호출", "Gemini 3.5 Flash-Lite completed 확인", "모델 ID와 minimal thinking 설정 동작"],
            ["Google Search 실호출", "최근 실행에서 HTTP 429 할당량 초과 재현", "오류 UI는 정상 동작, 할당량 복구 후 전체 E2E 재검증 필요"],
            ["보안 검사", "API 키 값과 URL query key 패턴 미포함", "키는 산출물에 포함하지 않음"],
        ], widths=[3.8, 4.8, 7.4])
    add_paragraph_before(
        doc, sect_pr,
        "최근 실제 브라우저 실행에서는 기사 분석까지 완료되었고 Google Search 단계에서 계정별 할당량 제한이 발생했습니다. 따라서 ‘전체 실시간 검증 완료’라고 과장하지 않고, 현재 상태를 quota 복구 후 재검증이 필요한 단계로 기록합니다.",
        kind="note",
    )

    add_paragraph_before(doc, sect_pr, "2. 추후 개선사항", kind="heading", keep_with_next=True)
    add_table_before(doc, sect_pr,
        ["우선순위", "개선 내용", "검증 방법"],
        [
            ["1", "429 재시도·대기시간·사용량 안내와 검색 결과 캐시", "할당량 초과·복구 시나리오 E2E"],
            ["2", "정부·공공기관·원문 등 출처 유형 표시와 복수 출처 교차 확인", "출처 품질 rubric과 오답 사례집"],
            ["3", "언론사별 추출 selector 회귀 테스트와 동적 페이지 대응", "사이트별 fixture·DOM 변화 테스트"],
            ["4", "색상 외 아이콘·밑줄·키보드 탐색·스크린리더 설명", "색각·키보드·접근성 점검"],
            ["5", "한국어 뉴스 검증 데이터셋으로 상태 정확도와 인용 적합성 평가", "사람이 라벨링한 표본과 비교"],
            ["6", "Interactions API Beta schema 변경 감시", "API revision별 contract test"],
        ], widths=[2.0, 8.6, 5.4])

    add_paragraph_before(doc, sect_pr, "■ 위험성 검토", kind="subheading", keep_with_next=True)
    add_bullets_before(doc, sect_pr, [
        "Fact Lens는 기사나 작성자를 자동 차단하지 않으며 법적·의학적·선거 관련 최종 판정을 대신하지 않습니다.",
        "AI가 잘못된 주장 추출이나 설명을 만들 수 있으므로 미확인 상태와 원문·출처 확인 절차를 유지합니다.",
        "기사 본문에 개인정보가 포함될 수 있으므로 민감 문서 사용을 피하고 API 전송 사실을 안내합니다.",
        "API 키는 비밀정보로 취급하고 화면·문서·로그에 값 자체를 출력하지 않습니다.",
        "색상만으로 상태를 구분하지 않도록 아이콘과 상태 문구를 함께 제공하고 접근성을 더 검증합니다.",
    ])

    # 출처는 공식 지침에 따라 마지막 붙임으로 추가합니다.
    page_break = add_paragraph_before(doc, sect_pr, "")
    page_break.add_run().add_break(WD_BREAK.PAGE)
    add_paragraph_before(doc, sect_pr, "붙임. 출처", kind="heading", keep_with_next=True)
    sources = [
        "1. Fact Lens 소스 코드, https://github.com/amanhasfallenintoriverincity/fact-lens (확인: 2026. 8. 16.)",
        "2. Google AI for Developers, Interactions API, https://ai.google.dev/gemini-api/docs/interactions",
        "3. Google AI for Developers, Grounding with Google Search, https://ai.google.dev/gemini-api/docs/interactions/google-search",
        "4. Google AI for Developers, API keys and security, https://ai.google.dev/gemini-api/docs/interactions/api-key",
        "5. Chrome for Developers, Extension service workers, https://developer.chrome.com/docs/extensions/develop/concepts/service-workers",
        "6. Chrome for Developers, Content scripts, https://developer.chrome.com/docs/extensions/develop/concepts/content-scripts",
        "7. Chrome Extensions, Manifest V3, https://developer.chrome.com/docs/extensions/develop/migrate/what-is-mv3",
        "8. React documentation, https://react.dev/",
        "9. Vite documentation, https://vite.dev/",
    ]
    for source in sources:
        p = add_paragraph_before(doc, sect_pr, source, kind="note")
        p.paragraph_format.left_indent = Cm(0.2)

    # 추가된 본문은 10~12pt, 검정색으로 제한합니다. 공식 안내문과 대단원 표는 원본을 유지합니다.
    doc.save(str(FORM2_OUT))


def write_checklist() -> None:
    checklist = """# Fact Lens 제출 전 확인사항

아래 항목은 임의로 추측하지 않고 비워 두었습니다. 제출자가 직접 확인해 작성해 주세요.

## 참가신청서

- [ ] 참가 구분 또는 부문
- [ ] 팀명
- [ ] 참가자 성명
- [ ] 생년월일 또는 학년·학번 등 공식 양식이 요구하는 정보
- [ ] 소속 학교
- [ ] 연락처와 이메일
- [ ] 보호자 정보와 동의 항목
- [ ] 개인정보 수집·이용 동의 체크
- [ ] 작성일
- [ ] 신청자·보호자 서명 또는 날인

## 서식 1 작품 요약서

- [ ] 팀명
- [ ] 참가자 1 성명·소속
- [ ] 참가자 2 성명·소속(해당 시)
- [ ] 참가자 3 성명·소속(해당 시)

## 서식 2 작품설명서

- [x] 작품 내용 작성
- [x] 계획 수립 단계 순수 SW 갈래 선택
- [x] Ⅰ~Ⅳ 대단원 작성
- [x] 출처 기재
- [x] 본문 15쪽 이내 확인(본문 8쪽)

## 제출 직전

- [ ] Word에서 DOCX 3종을 열어 팀 정보와 서명 최종 확인
- [ ] 참가 인원과 작품명이 대회 접수 시스템의 입력값과 일치하는지 확인
- [ ] Gemini/Google Search 할당량 복구 후 실제 뉴스 기사 E2E 재검증
- [ ] 확장 프로그램 `dist`를 Chrome에 다시 로드하고 뉴스 탭도 새로고침
- [ ] ZIP 파일명과 내부 파일명을 대회 공지의 제출 규칙에 맞게 변경
"""
    (OUTPUT_ROOT / "제출전_확인사항.md").write_text(checklist, encoding="utf-8")


def main() -> None:
    REQUIRED.mkdir(parents=True, exist_ok=True)
    WORK.mkdir(parents=True, exist_ok=True)
    fill_form1()
    fill_form2()
    shutil.copy2(APPLICATION_TEMPLATE, APPLICATION_OUT)
    write_checklist()
    print(FORM1_OUT)
    print(FORM2_OUT)
    print(APPLICATION_OUT)


if __name__ == "__main__":
    main()
